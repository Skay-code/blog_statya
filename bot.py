
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# Установка зависимостей
!pip install python-docx docxcompose beautifulsoup4 ebooklib aiogram aiofiles nest_asyncio mammoth
import mammoth
import os
import re
import time
import docx
import aiogram
from docx import Document
from docx.shared import Inches
from PIL import Image
import io
from io import BytesIO
import base64
import posixpath
from docxcompose.composer import Composer
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
from aiogram import Bot, Router, types, F, Dispatcher
from aiogram.types import Message, FSInputFile, BotCommand, BotCommandScopeDefault, BotCommandScopeAllGroupChats
from aiogram.filters import Command
from aiogram.utils.keyboard import ReplyKeyboardBuilder
from aiogram.utils import markdown as md
import aiofiles
import asyncio
import nest_asyncio
import concurrent.futures
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State
from aiogram.fsm.state import StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from functools import partial
from collections import deque
from datetime import datetime, timezone, timedelta
nest_asyncio.apply()
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery, ReplyKeyboardRemove
from aiogram.exceptions import TelegramBadRequest

# Создаем пул потоков для выполнения CPU-bound задач
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=1)

async def set_bot_commands(bot: Bot):
    commands = [
        BotCommand(command="start_merge", description="Начать сбор файлов"),
        BotCommand(command="end_merge", description="Завершить сбор и объединить"),
        BotCommand(command="cancel", description="Отменить сбор"),
        BotCommand(command="queue_status", description="Статус очереди задач"),
        BotCommand(command="limits", description="Проверить лимиты"),
        BotCommand(command="info", description="Информация о боте и команды"),
    ]
    # Команды для личных чатов
    await bot.set_my_commands(commands, scope=BotCommandScopeDefault())
    # Команды для всех групповых чатов
    await bot.set_my_commands(commands, scope=BotCommandScopeAllGroupChats())

async def sanitize_filename(file_name):
    replacement = '_'
    invalid_chars_pattern = r'[<>:"/|\?*]' # Паттерн для поиска недопустимых символов

    # Заменяем все недопустимые символы на specified replacement
    sanitized = re.sub(invalid_chars_pattern, replacement, file_name)

    # Опционально: ограничить длину имени файла, т.к. у ФС есть лимиты (обычно 255 байт)
    max_len = 250
    sanitized = sanitized[:max_len]

    return sanitized

async def check_sender(message: types.Message):
    """Проверяет отправителя. Если не пользователь, отвечает и возвращает True."""
    if message.sender_chat:
        bot_message = await message.reply("Анонимные пользователи (от имени каналов/групп) не могут использовать этого бота.")
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        return True # Да, это не пользователь, обработку надо прервать
    return False # Нет, это пользователь, можно продолжать

# --- Вспомогательная функция для удаления сообщений ---
async def delete_message_after_delay(message: types.Message, delay: int):
    """Удаляет сообщение после указанной задержки."""
    await asyncio.sleep(delay)
    try:
        await message.delete()
    except TelegramBadRequest: # Игнорируем ошибки, если сообщение уже удалено
        pass
    except Exception as e:
        print(f"Не удалось удалить сообщение {message.message_id}: {e}")

async def del_msg(chat_id, list_delete_message):
    # Удаляем сохранённые сообщения
    for msg_id in list_delete_message:
        try:
           await bot.delete_message(chat_id, msg_id)
           await asyncio.sleep(0.1)
        except TelegramBadRequest: pass
        except Exception as e: print(f"Ошибка удаления сообщения {msg_id} при end_merge: {e}")

class UserLimits:
    def __init__(self, max_files, max_size):
        self.user_data = {}  # {user_id: {'files_today': int}}
        self.last_global_reset = self._get_last_utc_midnight()
        self.user_locks = {} # Словарь для хранения блокировок пользователей
        self.max_files = max_files
        self.max_size = max_size
        self.admins = [5787446293, 5491435817]

    def _get_last_utc_midnight(self):
        """Возвращает последнюю полночь по UTC."""
        now = datetime.now(timezone.utc)
        return now.replace(hour=0, minute=0, second=0, microsecond=0)

    def get_lock(self, user_id):
        """Получает или создает блокировку для пользователя."""
        if user_id not in self.user_locks:
            self.user_locks[user_id] = asyncio.Lock()
        return self.user_locks[user_id]

    def check_limits(self, user_id, file_size):
        """Проверяет лимиты и сбрасывает их в 00:00 UTC."""
        now = datetime.now(timezone.utc)

        # Если наступил новый день (00:00 UTC), сбрасываем счетчики у всех
        if now > self.last_global_reset + timedelta(days=1):
            self.user_data.clear()  # Обнуляем данные всех пользователей
            self.last_global_reset = self._get_last_utc_midnight()

        # Инициализируем данные пользователя, если их нет
        if user_id not in self.user_data:
            self.user_data[user_id] = {'files_today': 0}

        # Проверяем лимиты
        if file_size > self.max_size * 1024 * 1024:  # Допустимый размер файла
            return False, f"❌ Размер файла превышает {self.max_size} MB."

        if user_id in self.admins:
            return True, ""

        if self.user_data[user_id]['files_today'] == self.max_files:
            time_left = (self.last_global_reset + timedelta(days=1)) - now
            hours_left = time_left.seconds // 3600
            minutes_left = (time_left.seconds % 3600) // 60
            return False, f"❌ Лимит исчерпан ({self.max_files}/{self.max_files}). Сброс через {hours_left} ч. {minutes_left} мин. (в 00:00 UTC)."

        return True, ""
    def increment_counter(self, user_id):
        """Увеличивает счетчик файлов пользователя."""
        self.user_data[user_id]['files_today'] += 1

    def discrement_counter(self, user_id, count):
        if user_id in self.user_data:
            self.user_data[user_id]['files_today'] -= count

# Создаем экземпляр класса лимитов
user_limits = UserLimits(max_files=30, max_size=15)

# Система очереди
class TaskQueue:
    def __init__(self, max_concurrent_tasks):
        self.queue = deque()  # Очередь задач
        self.active_tasks = {}  # Активные задачи: task_id -> task (вместо user_id -> task)
        self.max_concurrent_tasks = max_concurrent_tasks
        self.task_counter = 0  # Счетчик задач для назначения номера очереди

    def add_task(self, user_id, chat_id, message_thread_id, is_forum, file_list, output_file_name):
        """Добавляет задачу в очередь и возвращает уникальный ID задачи и позицию в очереди"""
        self.task_counter += 1
        task_id = self.task_counter
        task = {
            'user_id': user_id,
            'chat_id': chat_id,
            'message_thread_id': message_thread_id,
            'is_forum': is_forum,
            'file_list': file_list,
            'output_file_name': output_file_name,
            'task_id': task_id,
            'time_added': time.time(),
            'list_delete_message': []
        }
        self.queue.append(task)

        return task, len(self.queue)

    def get_next_task(self):
        """Получить следующую задачу из очереди"""
        if not self.queue:
            return None
        task = self.queue.popleft()
        self.active_tasks[task['task_id']] = task  # Используем task_id вместо user_id
        return task

    def complete_task(self, task_id):
        """Пометить задачу как завершенную"""
        if task_id in self.active_tasks:
            task = self.active_tasks[task_id]
            user_id = task['user_id']

            # Удаляем задачу из active_tasks
            del self.active_tasks[task_id]

    def get_user_tasks(self, user_id):
        """Получить список всех задач пользователя (в очереди и активных)"""
        tasks = []
        # Ищем в активных задачах
        for task_id, task in self.active_tasks.items():
            if task['user_id'] == user_id:
                tasks.append(task)

        # Ищем в очереди
        for task in self.queue:
            if task['user_id'] == user_id:
                tasks.append(task)

        return tasks

    def can_process_now(self):
        """Проверка, можно ли обработать следующую задачу из очереди"""
        return len(self.active_tasks) < self.max_concurrent_tasks and self.queue

# Создаем очередь задач
task_queue = TaskQueue(max_concurrent_tasks=1)  # Максимум 5 одновременных задач

# Декоратор для измерения времени выполнения функции
def timer(func):
    async def wrapper(*args, **kwargs):
        start_time = time.time()
        result = await func(*args, **kwargs)
        elapsed = time.time() - start_time
        print(f"[PROFILING] Функция {func.__name__} выполнилась за {elapsed:.2f} секунд")
        return result
    return wrapper

# Замените токен на свой
API_TOKEN = ''
bot = Bot(token=API_TOKEN)
router = Router()

# ===================== Неблокирующие функции конвертации =====================

# Функция-обертка для выполнения блокирующих операций в отдельном потоке
async def run_in_threadpool(func, *args, **kwargs):
    loop = asyncio.get_running_loop()
    func_partial = partial(func, *args, **kwargs)
    return await loop.run_in_executor(thread_pool, func_partial)

async def convert_docx_to_html_mammoth(docx_file_path):
    """Асинхронно конвертирует DOCX файл в HTML строку с помощью mammoth."""
    def _convert():
        try:
            with open(docx_file_path, "rb") as docx_file_obj:
                result = mammoth.convert_to_html(docx_file_obj)
                return result.value
        except Exception as e:
            print(f"Ошибка конвертации DOCX '{docx_file_path}' в HTML: {e}")
            return None # Возвращаем None в случае ошибки
    return await run_in_threadpool(_convert)

# Неблокирующие версии функций конвертации
async def convert_epub_to_docx(epub_file, docx_file):
    def _convert():
        document = Document()
        try:
            # Открываем EPUB-файл
            book = epub.read_epub(epub_file)
            spine_ids = [item[0] for item in book.spine]  # [ 'titlepage', 'Section0001.html', ... ]
            # Перебираем элементы книги
            for id_ in spine_ids:
                item = book.get_item_with_id(id_)
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.content, 'html.parser')
                    html_base_path = posixpath.dirname(item.get_name())
                    for element in soup.find_all():
                        if element.name == 'h1':
                            document.add_heading(element.get_text(), level=1)
                        elif element.name == 'p':
                            doc_paragraph = document.add_paragraph()
                            # Перебор вложенных элементов абзаца
                            for sub in element.contents:
                                if hasattr(sub, 'name'):
                                    if sub.name == 'strong':
                                        run = doc_paragraph.add_run(sub.get_text())
                                        run.bold = True
                                    elif sub.name == 'em':
                                        run = doc_paragraph.add_run(sub.get_text())
                                        run.italic = True
                                    else:
                                        doc_paragraph.add_run(sub.get_text())
                                else:
                                    # Если это просто текст
                                    doc_paragraph.add_run(sub)
                        # Обработка тегов изображений <img>
                        elif element.name == 'img':
                            src = element.get('src')
                            if src:
                                try:
                                    # Формируем полный путь к изображению внутри EPUB
                                    image_href = posixpath.normpath(posixpath.join(html_base_path, src))
                                    # Ищем элемент изображения в книге по его пути (href)
                                    img_item = book.get_item_with_href(image_href)

                                    if img_item and img_item.get_type() == ebooklib.ITEM_IMAGE:
                                        # Получаем бинарные данные изображения
                                        image_data = img_item.content
                                        try:
                                            # Оборачиваем данные в BytesIO, чтобы python-docx мог их прочитать
                                            image_stream = io.BytesIO(image_data)
                                            # Добавляем изображение в документ
                                            document.add_picture(image_stream, width=Inches(5.5))
                                        except:
                                            try:
                                                f = io.BytesIO()
                                                Image.open(io.BytesIO(image_data)).convert('RGB').save(f, format='JPEG')
                                                document.add_picture(f, width=Inches(5.5))
                                            except Exception as img_e:
                                                print(f"FB2: Ошибка добавления изображения '{image_id_ref}' в DOCX: {img_e}")
                                                document.add_paragraph(f"[Ошибка добавления изображения: {image_id_ref}]")
                                    else:
                                        print(f"Предупреждение: Не найден элемент изображения или тип не ITEM_IMAGE для href: {image_href} (src: {src})")
                                except KeyError:
                                    # Если get_item_with_href не нашел элемент
                                     print(f"Предупреждение: Не найден элемент изображения в манифесте EPUB для href: {image_href} (src: {src}) в файле {item.get_name()}")
                                except Exception as img_e:
                                    print(f"Ошибка при обработке изображения {src} в файле {item.get_name()}: {img_e}")
        except Exception as e:
            print(f"Ошибка конвертации EPUB {epub_file}: {e}")
            document.add_paragraph(f"Ошибка конвертации файла {os.path.basename(epub_file)}: {e}")
        finally:
            document.save(docx_file)

    return await run_in_threadpool(_convert)

async def convert_fb2_to_docx(fb2_file, docx_file):
    def _convert():
        document = Document()
        image_data_map = {} # Словарь для раскодированных изображений {id: image_bytes}
        try:
            with open(fb2_file, 'r', encoding='utf-8') as f:
                content = f.read()
            soup = BeautifulSoup(content, 'xml')
            # Извлечение и декодирование всех изображений
            for binary_tag in soup.find_all('binary'):
                image_id = binary_tag.get('id')
                content_type = binary_tag.get('content-type', '')
                data = binary_tag.text.strip()
                if image_id and data and content_type.startswith('image/'):
                    try:
                        image_bytes = base64.b64decode(data)
                        image_data_map[image_id] = image_bytes
                    except Exception as b64e:
                        print(f"FB2: Ошибка декодирования base64 для ID '{image_id}': {b64e}")
            # Парсим остальные части документа
            for element in soup.find_all(['title', 'p', 'image']):
                if element.name == 'title':
                    document.add_heading(element.get_text(), level=1)
                elif element.name == 'p':
                    # Если абзац не является частью title или annotation
                    if element.find_parent(['title', 'annotation']) is None:
                        doc_paragraph = document.add_paragraph()
                        for sub in element.contents:
                            if hasattr(sub, 'name'):
                                if sub.name == 'strong':
                                    run = doc_paragraph.add_run(sub.get_text())
                                    run.bold = True
                                elif sub.name == 'emphasis':
                                    run = doc_paragraph.add_run(sub.get_text())
                                    run.italic = True
                                else:
                                    doc_paragraph.add_run(sub.get_text())
                            else:
                                doc_paragraph.add_run(sub)
                # Обработка тега image
                elif element.name == 'image':
                    href_attr = element.get('l:href') or element.get('xlink:href')
                    if href_attr and href_attr.startswith('#'):
                        image_id_ref = href_attr[1:]
                        if image_id_ref in image_data_map:
                            image_bytes = image_data_map[image_id_ref]
                            try:
                                image_stream = io.BytesIO(image_bytes)
                                document.add_picture(image_stream, width=Inches(5.5))
                            except:
                                try:
                                    f = io.BytesIO()
                                    Image.open(io.BytesIO(image_bytes)).convert('RGB').save(f, format='JPEG')
                                    # f.seek(0)
                                    document.add_picture(f, width=Inches(5.5))
                                except Exception as img_e:
                                    print(f"FB2: Ошибка добавления изображения '{image_id_ref}' в DOCX: {img_e}")
                                    document.add_paragraph(f"[Ошибка добавления изображения: {image_id_ref}]")
                        else:
                            print(f"FB2: Данные для изображения '{image_id_ref}' не найдены.")
                            document.add_paragraph(f"[Изображение не найдено: {image_id_ref}]")
                    else:
                        print(f"FB2: Тег <image> без корректной ссылки: {element}")
                        document.add_paragraph("[Некорректный тег image]")
        except Exception as e:
            print(f"Ошибка конвертации FB2 {fb2_file}: {e}")
            document.add_paragraph(f"Ошибка конвертации файла {os.path.basename(fb2_file)}: {e}")
        finally:
            document.save(docx_file)
    return await run_in_threadpool(_convert)

async def convert_txt_to_docx(txt_file, docx_file):
    def _convert():
        try:
            with open(txt_file, 'r', encoding='utf-8') as f:
                text = f.read()
            document = Document()
            for line in text.splitlines():
                document.add_paragraph(line)
        except Exception as e:
            print(f"Ошибка конвертации TXT {txt_file}: {e}")
            # Создаем пустой docx или с сообщением об ошибке, чтобы процесс не падал
            document = Document()
            document.add_paragraph(f"Ошибка конвертации файла {os.path.basename(txt_file)}: {e}")
        finally:
            document.save(docx_file)

    return await run_in_threadpool(_convert)

@timer
async def process_files(file_list):
    """
    Обрабатывает список файлов, конвертируя их в формат .docx (если требуется)
    и возвращает список имен файлов в формате .docx для последующего объединения.
    """
    converted_files = []
    for file in file_list:
        ext = os.path.splitext(file)[1].lower()
        # Если файл уже в формате .docx – добавляем его в список
        if ext == ".docx":
            converted_files.append(file)
        elif ext == ".txt":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_txt_to_docx(file, docx_file)
            converted_files.append(docx_file)
        elif ext == ".fb2":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_fb2_to_docx(file, docx_file)
            converted_files.append(docx_file)
        elif ext == ".epub":
            docx_file = os.path.splitext(file)[0] + ".docx"
            await convert_epub_to_docx(file, docx_file)
            converted_files.append(docx_file)
    return converted_files

# ===================== Неблокирующие функции для работы с документами =====
def safe_docx(doc):
    check = Document()
    composer = Composer(check)
    composer.append(doc)
    return check

def check_and_add_title(doc, file_name):
    """
    Проверяет первые абзацы документа на наличие заголовка (например, "Глава ...").
    Если заголовок не найден, добавляет его на основе имени файла.
    """
    patterns = [
        r'Глава[ ]{0,4}\d{1,4}',
        r'Часть[ ]{0,4}\d{1,4}',
        r'^Пролог[ .!]*$',
        r'^Описание[ .!]*$',
        r'^Аннотация[ .!]*$',
        r'^Annotation[ .!]*$',
        r'^Предисловие от автора[ .!]*$'
    ]
    if doc.paragraphs:
        check_paragraphs = doc.paragraphs[0:4]
        title_found = False
        c = 0
        for p in check_paragraphs:
            if any(p.style.name.lower().startswith(prefix) for prefix in ["heading", "заголовок"]):
                title_found = True
                break

        if not title_found:
            for p in check_paragraphs:
                for pattern in patterns:
                    if re.search(pattern, p.text.strip()):
                        title_found = True
                        try:
                            p.style = 'Heading 1'
                        except Exception as e:
                            try:
                                doc = safe_docx(doc)
                                p = doc.paragraphs[c]
                                p.style = 'Heading 1'
                            except Exception as e:
                                print(f"Возникла ошибка при создании заголовка: {e}")
                        break
                if title_found:
                    break
                c = c+1

        if not title_found:
            # Добавляем заголовок перед первым абзацем
            title = os.path.splitext(os.path.basename(file_name))[0]
            if re.fullmatch(r'\d+', title.strip()):
                title = 'Глава ' + title
            try:
                paragraph = doc.paragraphs[0].insert_paragraph_before(title)
                paragraph.style = 'Heading 1'
            except:
                try:
                    doc = safe_docx(doc)
                    paragraph = doc.paragraphs[0]
                    paragraph.style = 'Heading 1'
                    return doc
                except Exception as e:
                    print(f"Возникла ошибка при добавлении заголовка: {e}")
    return doc

@timer
async def merge_docx(file_list, output_file_name):
    def _merge():
        # Создаем новый документ
        merged_document = Document()
        composer = Composer(merged_document)
        try:
            for file in file_list:
                try:
                    doc = Document(file)
                    doc = check_and_add_title(doc, file)
                    composer.append(doc)
                except Exception as e:
                    print(f"Ошибка добавления файла {file}: {e}")
                    merged_document.add_paragraph(f"Ошибка добавления файла {os.path.basename(file)}: {e}")
        except Exception as e:
            print(f"Критическая ошибка, невозможно пройтись по списку {file_list}: {e}")
            merged_document.add_paragraph(f"Критическая ошибка, невозможно пройтись по списку {file_list}: {e}")
        finally:
            composer.save(output_file_name)
            print(f"Файлы объединены в {output_file_name}")
            return output_file_name

    # Объединяем обработанные файлы в отдельном потоке
    result = await run_in_threadpool(_merge)
    return result

# ===================== FSM: Состояния =====================
class MergeStates(StatesGroup):
    collecting = State()  # Состояние сбора файлов
    naming_file = State() # Состояние запроса имени файла

# ===================== Обработчики Telegram-бота =====================
@router.message(Command("start_merge"))
async def start_merge(message: Message, state: FSMContext):
    """
    Команда для начала сбора файлов.
    """
    if await check_sender(message):
        return

    current_state = await state.get_state()
    if current_state == MergeStates.collecting.state:
        bot_message = await message.answer("Сбор файлов уже запущен.")
        await message.delete()
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        return

    # Теперь мы не проверяем, есть ли у пользователя активная задача
    # Просто начинаем новый сбор файлов

    await state.set_state(MergeStates.collecting)
    bot_message = await message.answer("Сбор файлов начат! Отправляйте файлы. Используйте /end_merge для завершения или /cancel для отмены.")
    await state.update_data(file_list=[], list_delete_message=[bot_message.message_id]) # Создаем пустой список файлов
    await message.delete()

def build_task_status(user_id):
    user_tasks = task_queue.get_user_tasks(user_id)
    if not user_tasks:
        total_tasks = len(task_queue.queue)
        active_tasks = len(task_queue.active_tasks)
        text = f"У вас нет задач в очереди.\nСтатус системы: {active_tasks}/{task_queue.max_concurrent_tasks} активных задач, {total_tasks} задач в очереди."
        return text, None

    # Формируем сообщение со списком задач пользователя
    tasks_info = []
    keyboard_buttons = []
    for task in user_tasks:
        task_id = task['task_id']

        # Проверяем, активна ли задача
        if task_id in task_queue.active_tasks:
            status = "⚙️ Выполняется (отменить невозможно)"
        else:
            # Ищем позицию в очереди
            for i, queued_task in enumerate(task_queue.queue):
                if queued_task['task_id'] == task_id:
                    status = f"🕒 В очереди (позиция {i+1})"
                    break

        # Создаем имя задачи из первого файла в списке
        task_name = os.path.basename(task['file_list'][0])
        if len(task['file_list']) > 1:
            task_name += f" и еще {len(task['file_list'])-1} файлов"

        tasks_info.append(f"Задача #{task_id}: {task_name} - {status}")
        if task_id not in task_queue.active_tasks:
            # Добавляем кнопку для этой задачи, если он не выполняется
            keyboard_buttons.append(
               InlineKeyboardButton(text=f"Отменить #{task_id}", callback_data=f"cancel:{task_id}")
            )

    text = "Ваши задачи:\n\n" + "\n".join(tasks_info)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[keyboard_buttons[i:i+2] for i in range(0, len(keyboard_buttons), 2)])
    return text, keyboard

@router.message(Command("queue_status"))
async def queue_status(message: Message):
    """
    Проверка статуса очереди.
    """
    if await check_sender(message):
        return

    user_id = message.from_user.id
    text, keyboard = build_task_status(user_id)
    bot_message = await message.answer(text, reply_markup=keyboard)
    asyncio.create_task(delete_message_after_delay(bot_message, delay=300))
    await message.delete()

@router.message(Command("cancel"))
async def cancel_collecting(message: Message, state: FSMContext):
    """
    Отмена сбора файлов.
    """
    if await check_sender(message):
        return

    current_state = await state.get_state()
    if current_state != MergeStates.collecting.state:
        bot_message = await message.answer("Сбор файлов не был запущен.")
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        await message.delete()
        return

    # Получаем список файлов, чтобы удалить их
    user_data = await state.get_data()
    file_list = user_data.get('file_list', [])
    list_delete_message = user_data.get('list_delete_message', [])
    chat_id = message.chat.id
    user_id = message.from_user.id

    # Удаляем сохранённые сообщения
    await del_msg(chat_id, list_delete_message)
    user_limits.discrement_counter(user_id, len(file_list))
    max_files = user_limits.max_files
    if user_id in user_limits.user_data:
        files_today_count = user_limits.user_data[user_id]['files_today']
    else:
        files_today_count = 0

    # Удаляем временные файлы
    for file_item in file_list:
        file = file_item[0]
        if os.path.exists(file):
            os.remove(file)

    await state.clear()
    bot_message = await message.answer("Сбор файлов отменен. Все временные файлы удалены.\n"
                  f"Ваш лимит: {files_today_count}/{max_files} (-{len(file_list)})" # Показываем актуальное число
                  )
    asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
    await message.delete()

@router.callback_query(lambda c: c.data.startswith("cancel:"))
async def handle_cancel_callback(callback_query: CallbackQuery):
    user_id = callback_query.from_user.id
    task_id = int(callback_query.data.split(":")[1])
    message = callback_query.message

    # Ищем задачу в очереди
    found = False
    new_queue = deque()
    for task in task_queue.queue:
        if task['task_id'] == task_id:
            if task['user_id'] == user_id:
                found = True
                # Удаляем временные файлы
                for file in task['file_list']:
                    if os.path.exists(file):
                        os.remove(file)
            else:
                # Задача существует, но принадлежит другому пользователю
                await message.answer("Вы не можете отменить чужую задачу")
                return
        else:
            new_queue.append(task)

    if found:
        # Обновляем очередь
        task_queue.queue = new_queue
        text, keyboard = build_task_status(user_id)
        await message.edit_text(text, reply_markup=keyboard)

        file_list = task['file_list']
        user_limits.discrement_counter(user_id, len(file_list))
        max_files = user_limits.max_files
        files_today_count = user_limits.user_data[user_id]['files_today']
        bot_message = await message.answer("Задача #{task_id} удалена из очереди\n"
            f"Ваш лимит: {files_today_count}/{max_files} (-{len(file_list)})" # Показываем актуальное число
        )
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
    else:
        # Проверяем, не выполняется ли задача в данный момент
        if task_id in task_queue.active_tasks and task_queue.active_tasks[task_id]['user_id'] == user_id:
            await message.answer(f"Задача #{task_id} уже выполняется и не может быть отменена")
        else:
            await message.answer(f"Задача #{task_id} не найдена")

@router.message(Command("end_merge"))
async def end_merge(message: Message, state: FSMContext):
    """
    Команда для завершения сбора файлов и запроса имени выходного файла.
    """
    if await check_sender(message):
        return

    current_state = await state.get_state()
    if current_state != MergeStates.collecting.state:
        bot_message = await message.answer("Сбор файлов не был запущен. Введите /start_merge для начала.")
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        await message.delete()
        return

    user_data = await state.get_data()
    file_list = user_data.get('file_list', [])
    list_delete_message = user_data.get('list_delete_message', [])
    chat_id = message.chat.id

    if not file_list:
        bot_message = await message.answer("Нет файлов для обработки!")
        await state.clear()  # Очищаем состояние
        asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        await message.delete()
        # Удаляем сохранённые сообщения
        await del_msg(chat_id, list_delete_message)
        return

    # Переходим к состоянию запроса имени файла
    await state.set_state(MergeStates.naming_file)
    list_delete_message = user_data.get('list_delete_message', [])

    # Создаем клавиатуру с кнопкой "Пропустить"
    keyboard = ReplyKeyboardBuilder()
    keyboard.add(types.KeyboardButton(text="Пропустить"))
    keyboard.adjust(1)

    bot_message = await message.answer(
        "Введите название для итогового файла или нажмите 'Пропустить' для использования стандартного имени (merged.docx):",
        reply_markup=keyboard.as_markup(resize_keyboard=True)
    )
    list_delete_message.append(bot_message.message_id)
    await state.update_data(list_delete_message=list_delete_message)
    await message.delete()

@router.message(MergeStates.naming_file)
async def process_filename(message: Message, state: FSMContext):
    """
    Обработка введенного имени файла.
    """
    user_id = message.from_user.id
    chat_id = message.chat.id
    message_thread_id = message.message_thread_id
    is_forum = message.is_topic_message
    user_data = await state.get_data()
    file_list = user_data.get('file_list', [])
    list_delete_message = user_data.get('list_delete_message', [])

    # Сортируем файлы по ID сообщения (второй элемент кортежа)
    file_list.sort(key=lambda x: x[1])

    # Извлекаем только имена файлов после сортировки
    sorted_files = [file[0] for file in file_list]

    # Определяем имя выходного файла
    if message.text == "Пропустить":
        output_file_name = "merged.docx"
    else:
        output_file_name = message.text + ".docx"
        output_file_name = await sanitize_filename(output_file_name)

    # Добавляем задачу в очередь с отсортированным списком файлов
    task, queue_position = task_queue.add_task(user_id, chat_id, message_thread_id, is_forum, sorted_files, output_file_name)
    await message.delete()

    if queue_position > 0:
        bot_message = await message.answer(
            f"Итоговый файл будет назван: {output_file_name}\n"
            f"Ваша задача добавлена в очередь на позицию {queue_position}.\n"
            f"Используйте /queue_status для проверки статуса.", reply_markup=ReplyKeyboardRemove()
            )
        list_delete_message.append(bot_message.message_id)
        task['list_delete_message'] = list_delete_message

    # Очищаем состояние после добавления задачи в очередь
    await state.clear()

    # Пытаемся запустить обработку задачи, если есть свободные потоки
    asyncio.create_task(check_and_process_queue())

async def check_and_process_queue():
    """
    Проверяет очередь и запускает обработку новых задач, если есть свободные ресурсы.
    """
    while task_queue.can_process_now():
        task = task_queue.get_next_task()
        if task:
            chat_id = task['chat_id']
            message_thread_id = task['message_thread_id']
            is_forum = task['is_forum']
            file_list = task['file_list']
            output_file_name = task['output_file_name']
            task_id = task['task_id']
            list_delete_message = task['list_delete_message']

            send_kwargs = {} # Словарь для тем
            if is_forum == True: # Проверяем, существуют ли темы в группе
                    send_kwargs["message_thread_id"] = message_thread_id

            # Уведомляем пользователя о начале обработки
            bot_message = await bot.send_message(chat_id, f"Начинаю обработку задачи #{task_id} с {len(file_list)} файлами. Это может занять некоторое время...", **send_kwargs )
            list_delete_message.append(bot_message.message_id)
            # Запускаем обработку в фоновом режиме
            asyncio.create_task(process_and_merge_files_with_queue(chat_id, send_kwargs, file_list, list_delete_message, output_file_name, task_id))

async def process_and_merge_files_with_queue(chat_id, send_kwargs, file_list, list_delete_message, output_file_name, task_id):
    """
    Асинхронная функция для обработки и объединения файлов с учетом очереди.
    """
    try:
        # Конвертация и объединение файлов
        converted_files = await process_files(file_list)
        merged_file = await merge_docx(converted_files, output_file_name)

        # Формируем сообщение с информацией о собранных файлах
        file_list_str = "\n".join([os.path.basename(f) for f in file_list])
        await bot.send_message(chat_id, f"Задача #{task_id} завершена!\nФайлы объединены в {os.path.basename(output_file_name)}.\nСобрано {len(file_list)} файлов:\n{file_list_str}", **send_kwargs)

        # Отправляем объединённый файл пользователю
        document = FSInputFile(merged_file)
        caption = os.path.splitext(output_file_name)[0]
        await bot.send_document(chat_id, document=document, caption=caption, **send_kwargs)

        # Конвертируем объединенный DOCX в HTML
        status_msg_html = await bot.send_message(chat_id, f"Начинаю конвертацию {os.path.basename(merged_file)} в HTML...", **send_kwargs)
        list_delete_message.append(status_msg_html.message_id)

        html_content_from_merged_docx = await convert_docx_to_html_mammoth(merged_file)

        if html_content_from_merged_docx:
            output_html_filename = os.path.splitext(merged_file)[0] + ".html"
            try:
                async with aiofiles.open(output_html_filename, "w", encoding="utf-8") as f:
                    await f.write(html_content_from_merged_docx)

                # Отправляем HTML файл пользователю
                document_html = FSInputFile(output_html_filename)
                caption_html = os.path.splitext(os.path.basename(output_html_filename))[0]
                await bot.send_document(chat_id, document=document_html, caption=caption_html, **send_kwargs)
                await bot.send_message(chat_id, f"Файл {os.path.basename(output_html_filename)} успешно сконвертирован и отправлен.", **send_kwargs)

                # Удаляем временный HTML файл после отправки
                if os.path.exists(output_html_filename):
                    os.remove(output_html_filename)

            except Exception as e_html_save_send:
                await bot.send_message(chat_id, f"Ошибка при сохранении или отправке HTML файла: {e_html_save_send}", **send_kwargs)
        else:
            await bot.send_message(chat_id, f"Не удалось сконвертировать {os.path.basename(merged_docx_file_path)} в HTML.", **send_kwargs)

        # Удаляем объединенный файл после отправки
        if os.path.exists(merged_file):
            os.remove(merged_file)

    except Exception as e:
        await bot.send_message(chat_id, f"Произошла ошибка при обработке задачи #{task_id}: {str(e)}", **send_kwargs)

    finally:
        # Удаляем сохранённые сообщения
        await del_msg(chat_id, list_delete_message)

        # Удаляем файлы, отправленные пользователем
        for file in file_list:
            if os.path.exists(file):
                os.remove(file)

        # Отмечаем задачу как выполненную
        task_queue.complete_task(task_id)  # Теперь передаю task_id (раньше было user_id)

        # Проверяем, можно ли обработать следующую задачу
        asyncio.create_task(check_and_process_queue())

@router.message(F.document)
async def handle_document(message: Message, state: FSMContext):
    """
    Обработчик полученных файлов.
    Если сбор файлов запущен, сохраняет полученный документ на диск
    и добавляет его имя в список для дальнейшей обработки.
    """
    if await check_sender(message):
        return

    current_state = await state.get_state()
    if current_state != MergeStates.collecting.state:
        if message.chat.type == "private":
            bot_message = await message.answer("Сбор файлов не запущен. Введите /start_merge для начала.")
            asyncio.create_task(delete_message_after_delay(bot_message, delay=5))
        return

    file_name = message.document.file_name
    file_name = await sanitize_filename(file_name)
    base_name, extension = os.path.splitext(file_name)
    counter = 1

    if extension.lower() not in (".docx", ".fb2", ".txt", ".epub"):
        bot_message = await message.answer(f"Неизвестный формат файла: {message.document.file_name}. Пожалуйста, отправляйте файлы только в форматах docx, fb2, epub, txt.")
        asyncio.create_task(delete_message_after_delay(bot_message, delay=10))
        return

    user_id = message.from_user.id
    file_size = message.document.file_size
    lock = user_limits.get_lock(user_id) # Получаем блокировку пользователя

    async with lock: # Захватываем блокировку (освободится автоматически при выходе из блока)
        is_allowed, error_msg  = user_limits.check_limits(user_id, file_size)
        if not is_allowed:
            bot_message = await message.answer(error_msg)
            asyncio.create_task(delete_message_after_delay(bot_message, delay=10))
            return # Выходим, блокировка освобождается

        # Если лимит позволяет, СРАЗУ увеличиваем счетчик ВНУТРИ блокировки
        user_limits.increment_counter(user_id)
        max_files = user_limits.max_files
        files_today_count = user_limits.user_data[user_id]['files_today']

    # --- Операции вне блокировки (загрузка, сохранение) ---
    try:
        # Добавляем цифры к имени файла, если нужно, чтобы избежать конфликта между файлами
        while os.path.exists(file_name):
            file_name = f"{base_name}({counter}){extension}"
            counter += 1

        file_info = await bot.get_file(message.document.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)

        # Сохраняем файл на диск
        async with aiofiles.open(file_name, 'wb') as new_file:
            await new_file.write(downloaded_file.read())

        # Добавляем файл в список вместе с ID сообщения
        user_data = await state.get_data()
        file_list = user_data.get('file_list', [])
        list_delete_message = user_data.get('list_delete_message', [])
        # Теперь храним кортеж (имя_файла, id_сообщения)
        file_list.append((file_name, message.message_id))
        await state.update_data(file_list=file_list)

        # Сообщаем о лимитах
        bot_message = await message.answer(
            f"Файл {file_name} сохранён! Всего файлов: {len(file_list)}\n"
            f"Использовано сегодня: {files_today_count}/{max_files}" # Показываем актуальное число
        )
        list_delete_message.append(bot_message.message_id)
        await state.update_data(list_delete_message=list_delete_message)
    except Exception as e:
        await message.answer(f"Ошибка при сохранении файла: {str(e)}")

@router.message(Command("start"))
async def send_welcome(message: Message):
    if await check_sender(message):
        return
    await message.answer("Привет, я бот для объединения файлов! Нажми /info для получения дополнительной информации.")
    await message.delete()

@router.message(Command("info"))
async def send_info(message: Message):
    if await check_sender(message):
        return

    max_files = user_limits.max_files
    max_size = user_limits.max_size

    bot_message = await message.answer(
        "📚 Бот для объединения файлов (DOCX, FB2, EPUB, TXT).\n\n"
        "Лимиты:\n"
        f"• {max_files} файлов в сутки (сброс в 00:00 UTC)\n"
        f"• Макс. размер файла: {max_size} MB\n\n"
        "Команды:\n"
        "/start_merge – начать сбор файлов\n"
        "/end_merge – завершить и объединить\n"
        "/limits – проверить лимиты\n"
        "/queue_status – статус очереди\n"
        "/cancel – отменить текущий сбор"
    )
    asyncio.create_task(delete_message_after_delay(bot_message, delay=300))
    await message.delete()

@router.message(Command("limits"))
async def check_limits(message: Message):
    """Показывает текущие лимиты и время до сброса."""
    if await check_sender(message):
        return

    user_id = message.from_user.id
    now = datetime.now(timezone.utc)
    is_allowed, error_msg  = user_limits.check_limits(user_id, 0)
    next_reset = user_limits.last_global_reset + timedelta(days=1)
    time_left = next_reset - now
    hours_left = time_left.seconds // 3600
    minutes_left = (time_left.seconds % 3600) // 60

    max_files = user_limits.max_files
    max_size = user_limits.max_size
    files_used = user_limits.user_data[user_id]['files_today']
    files_left = max_files - files_used

    bot_message = await message.answer(
        f"📊 Ваши лимиты:\n"
        f"• Использовано файлов: {files_used}/{max_files}\n"
        f"• Осталось файлов: {files_left}\n"
        f"• Максимальный размер файла: {max_size} MB\n"
        f"Лимит сбросится в 00:00 UTC (через {hours_left} ч. {minutes_left} мин.)"
    )
    asyncio.create_task(delete_message_after_delay(bot_message, delay=300))
    await message.delete()

# ===================== Запуск бота =====================
async def main():
    storage = MemoryStorage()
    dp = Dispatcher(storage=storage)
    dp.include_router(router)
    await set_bot_commands(bot)
    print("Бот запущен.")
    await dp.start_polling(bot, allowed_updates=dp.resolve_used_update_types())

if __name__ == "__main__":
    asyncio.run(main())
